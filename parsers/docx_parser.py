from pathlib import Path

import docx


def parse_docx(file_path: str | Path) -> str:
    """Extract text from paragraphs and table cells of a DOCX file."""
    doc = docx.Document(str(file_path))
    chunks: list[str] = []

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            chunks.append(paragraph.text.strip())

    for table in doc.tables:
        for row in table.rows:
            row_values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_values:
                chunks.append(" | ".join(row_values))

    return "\n".join(chunks).strip()